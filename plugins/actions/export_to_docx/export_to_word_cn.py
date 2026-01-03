"""
title: 导出为 Word
author: Fu-Jie
author_url: https://github.com/Fu-Jie
funding_url: https://github.com/Fu-Jie/awesome-openwebui
version: 0.2.0
icon_url: data:image/svg+xml;base64,PHN2ZwogIHhtbG5zPSJodHRwOi8vd3d3LnczLm9yZy8yMDAwL3N2ZyIKICB3aWR0aD0iMjQiCiAgaGVpZ2h0PSIyNCIKICB2aWV3Qm94PSIwIDAgMjQgMjQiCiAgZmlsbD0ibm9uZSIKICBzdHJva2U9ImN1cnJlbnRDb2xvciIKICBzdHJva2Utd2lkdGg9IjIiCiAgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIgogIHN0cm9rZS1saW5lam9pbj0icm91bmQiCj4KICA8cGF0aCBkPSJNNiAyMmEyIDIgMCAwIDEtMi0yVjRhMiAyIDAgMCAxIDItMmg4YTIuNCAyLjQgMCAwIDEgMS43MDQuNzA2bDMuNTg4IDMuNTg4QTIuNCAyLjQgMCAwIDEgMjAgOHYxMmEyIDIgMCAwIDEtMiAyeiIgLz4KICA8cGF0aCBkPSJNMTQgMnY1YTEgMSAwIDAgMCAxIDFoNSIgLz4KICA8cGF0aCBkPSJNMTAgOUg4IiAvPgogIDxwYXRoIGQ9Ik0xNiAxM0g4IiAvPgogIDxwYXRoIGQ9Ik0xNiAxN0g4IiAvPgo8L3N2Zz4K
requirements: python-docx==1.1.2, Pygments>=2.15.0, latex2mathml, mathml2omml
description: 将对话导出为 Word (.docx)，支持代码高亮、原生数学公式 (LaTeX)、Mermaid 图表、引用参考和增强表格格式。
"""

import os
import re
import base64
import datetime
import io
import asyncio
import logging
from typing import (
    Optional,
    Callable,
    Awaitable,
    Any,
    List,
    Tuple,
    Union,
    Dict,
    Literal,
    cast,
)
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from open_webui.models.chats import Chats
from open_webui.models.users import Users
from open_webui.utils.chat import generate_chat_completion
from pydantic import BaseModel, Field
from dataclasses import dataclass


# Pygments for syntax highlighting
try:
    from pygments import lex
    from pygments.lexers import get_lexer_by_name, TextLexer
    from pygments.token import Token

    PYGMENTS_AVAILABLE = True
except ImportError:
    PYGMENTS_AVAILABLE = False

try:
    from latex2mathml.converter import convert as latex_to_mathml
    import mathml2omml

    LATEX_MATH_AVAILABLE = True
except Exception:
    LATEX_MATH_AVAILABLE = False

_REASONING_DETAILS_RE = re.compile(
    r"<details\b[^>]*\btype\s*=\s*(?:\"reasoning\"|'reasoning'|reasoning)[^>]*>.*?</details\s*>",
    re.IGNORECASE | re.DOTALL,
)
_THINK_RE = re.compile(r"<think\b[^>]*>.*?</think\s*>", re.IGNORECASE | re.DOTALL)
_ANALYSIS_RE = re.compile(
    r"<analysis\b[^>]*>.*?</analysis\s*>", re.IGNORECASE | re.DOTALL
)


logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class _CitationRef:
    idx: int
    anchor: str
    title: str
    url: Optional[str]
    source_id: str


@dataclass
class _MermaidFenceBlock:
    info_raw: str
    language: str
    attrs: List[str]
    source: str


class Action:
    class Valves(BaseModel):
        TITLE_SOURCE: str = Field(
            default="chat_title",
            description="标题来源: 'chat_title' (对话标题), 'ai_generated' (AI 生成), 'markdown_title' (Markdown 标题)",
        )
        MERMAID_JS_URL: str = Field(
            default="https://cdn.jsdelivr.net/npm/mermaid@10.9.1/dist/mermaid.min.js",
            description="Mermaid JS CDN URL",
        )
        MERMAID_JSZIP_URL: str = Field(
            default="https://cdnjs.cloudflare.com/ajax/libs/jszip/3.10.1/jszip.min.js",
            description="JSZip CDN URL (用于 DOCX 操作)",
        )
        MERMAID_OPTIMIZE_LAYOUT: bool = Field(
            default=True,
            description="优化 Mermaid 布局: 自动将 LR (左右) 转换为 TD (上下) 以适应页面。",
        )
        MERMAID_PNG_SCALE: float = Field(
            default=3.0,
            description="Mermaid PNG 缩放比例 (分辨率): 越高越清晰但文件越大。默认: 3.0",
        )
        MERMAID_DISPLAY_SCALE: float = Field(
            default=1.5,
            description="Mermaid 显示比例 (视觉大小): >1.0 放大, <1.0 缩小。默认: 1.5",
        )
        MERMAID_CAPTIONS_ENABLE: bool = Field(
            default=True,
            description="启用 Mermaid 图表标题",
        )
        MERMAID_CAPTION_STYLE: str = Field(
            default="Caption",
            description="Mermaid 标题样式名称",
        )
        MERMAID_CAPTION_PREFIX: str = Field(
            default="图",
            description="Mermaid 标题前缀",
        )

    def __init__(self):
        self.valves = self.Valves()
        self._mermaid_figure_counter = 0
        self._caption_style_name = ""
        self._citation_anchor_by_index: Dict[int, str] = {}
        self._citation_refs: List[_CitationRef] = []
        self._bookmark_id_counter: int = 1

    async def _send_notification(self, emitter: Callable, type: str, content: str):
        await emitter(
            {"type": "notification", "data": {"type": type, "content": content}}
        )

    async def action(
        self,
        body: dict,
        __user__=None,
        __event_emitter__=None,
        __event_call__: Optional[Callable[[Any], Awaitable[None]]] = None,
        __metadata__: Optional[dict] = None,
        __request__: Optional[Any] = None,
    ):
        logger.info(f"action:{__name__}")

        # 解析用户信息
        if isinstance(__user__, (list, tuple)):
            user_language = (
                __user__[0].get("language", "zh-CN") if __user__ else "zh-CN"
            )
            user_name = __user__[0].get("name", "用户") if __user__[0] else "用户"
            user_id = (
                __user__[0]["id"]
                if __user__ and "id" in __user__[0]
                else "unknown_user"
            )
        elif isinstance(__user__, dict):
            user_language = __user__.get("language", "zh-CN")
            user_name = __user__.get("name", "用户")
            user_id = __user__.get("id", "unknown_user")

        if __event_emitter__:
            last_assistant_message = body["messages"][-1]

            await __event_emitter__(
                {
                    "type": "status",
                    "data": {"description": "正在转换为 Word 文档...", "done": False},
                }
            )

            try:
                message_content = last_assistant_message["content"]

                if not message_content or not message_content.strip():
                    await self._send_notification(
                        __event_emitter__, "error", "没有找到可导出的内容！"
                    )
                    return

                # 生成文件名
                title = ""
                chat_id = self.extract_chat_id(body, __metadata__)

                # 直接通过 chat_id 获取标题，因为 body 中通常不包含标题
                chat_title = ""
                if chat_id:
                    chat_title = await self.fetch_chat_title(chat_id, user_id)

                # 根据配置决定文件名使用的标题
                if (
                    self.valves.TITLE_SOURCE == "chat_title"
                    or not self.valves.TITLE_SOURCE
                ):
                    title = chat_title
                elif self.valves.TITLE_SOURCE == "markdown_title":
                    title = self.extract_title(message_content)
                elif self.valves.TITLE_SOURCE == "ai_generated":
                    title = await self.generate_title_using_ai(
                        body, message_content, user_id, __request__
                    )

                current_datetime = datetime.datetime.now()
                formatted_date = current_datetime.strftime("%Y%m%d")

                if title:
                    filename = f"{self.clean_filename(title)}.docx"
                else:
                    filename = f"{user_name}_{formatted_date}.docx"

                # 创建 Word 文档；若正文无一级标题，使用对话标题作为一级标题
                # 如果选择了 chat_title 且获取到了，则作为 top_heading
                # 如果选择了其他方式，title 就是文件名，也可以作为 top_heading

                # 保持原有逻辑：top_heading 主要是为了在文档开头补充标题
                # 这里我们尽量使用 chat_title 作为 top_heading，如果它存在的话，因为它通常是对话的主题
                # 即使文件名是 AI 生成的，文档内的标题用 chat_title 也是合理的
                # 但如果用户选择了 markdown_title，可能不希望插入 chat_title

                top_heading = ""
                if chat_title:
                    top_heading = chat_title
                elif title:
                    top_heading = title

                has_h1 = bool(re.search(r"^#\s+.+$", message_content, re.MULTILINE))

                # Extract sources if available (for citations)
                sources = (
                    last_assistant_message.get("sources") or body.get("sources") or []
                )

                doc = self.markdown_to_docx(
                    message_content,
                    top_heading=top_heading,
                    has_h1=has_h1,
                    sources=sources,
                )

                # 保存到内存
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                file_content = doc_buffer.read()
                base64_blob = base64.b64encode(file_content).decode("utf-8")

                # Escape message_content for JavaScript template literal
                escaped_content = (
                    message_content.replace("\\", "\\\\")  # Escape backslashes first
                    .replace("`", "\\`")  # Escape backticks
                    .replace("${", "\\${")  # Escape template literal expressions
                )

                # 触发文件下载
                if __event_call__:
                    await __event_call__(
                        {
                            "type": "execute",
                            "data": {
                                "code": f"""
                                (async function() {{
                                    try {{
                                        // Parse document.xml to find placeholders and extract optimized code
                                        // We do this FIRST to get the actual code to render (which might have been optimized in Python)
                                        
                                        // Load JSZip
                                        if (!window.JSZip) {{
                                            await new Promise((resolve, reject) => {{
                                                const script = document.createElement("script");
                                                script.src = "{self.valves.MERMAID_JSZIP_URL}";
                                                script.onload = resolve;
                                                script.onerror = reject;
                                                document.head.appendChild(script);
                                            }});
                                        }}

                                        const base64Data = "{base64_blob}";
                                        const binaryData = atob(base64Data);
                                        const arrayBuffer = new Uint8Array(binaryData.length);
                                        for (let i = 0; i < binaryData.length; i++) {{
                                            arrayBuffer[i] = binaryData.charCodeAt(i);
                                        }}
                                        
                                        const zip = new JSZip();
                                        await zip.loadAsync(arrayBuffer);

                                        // Parse document.xml
                                        const docXml = await zip.file("word/document.xml").async("string");
                                        const parser = new DOMParser();
                                        const xmlDoc = parser.parseFromString(docXml, "application/xml");
                                        
                                        const drawings = xmlDoc.getElementsByTagName("w:drawing");
                                        const placeholderInfo = [];
                                        
                                        for (let i = 0; i < drawings.length; i++) {{
                                            const drawing = drawings[i];
                                            const docPr = drawing.getElementsByTagName("wp:docPr")[0];
                                            if (docPr) {{
                                                const descr = docPr.getAttribute("descr");
                                                if (descr && descr.startsWith("MERMAID_SRC:")) {{
                                                    const encodedCode = descr.substring("MERMAID_SRC:".length);
                                                    const code = decodeURIComponent(encodedCode);
                                                    
                                                    // Find the blip and extent to replace
                                                    const parent = drawing.parentNode; // w:r usually, or w:drawing parent
                                                    // We need to find a:blip and wp:extent within this drawing
                                                    const blip = drawing.getElementsByTagName("a:blip")[0];
                                                    const extent = drawing.getElementsByTagName("wp:extent")[0];
                                                    
                                                    if (blip && extent) {{
                                                        const rId = blip.getAttribute("r:embed");
                                                        placeholderInfo.push({{ rId, extent, code }});
                                                    }}
                                                }}
                                            }}
                                        }}

                                        if (placeholderInfo.length === 0) {{
                                            console.log("No Mermaid placeholders found in DOCX.");
                                            // Just download the file as is
                                            const blob = new Blob([arrayBuffer], {{type: "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}});
                                            const url = URL.createObjectURL(blob);
                                            const a = document.createElement("a");
                                            a.style.display = "none";
                                            a.href = url;
                                            a.download = "{filename}";
                                            document.body.appendChild(a);
                                            a.click();
                                            URL.revokeObjectURL(url);
                                            document.body.removeChild(a);
                                            return;
                                        }}
                                        
                                        console.log(`Found ${{placeholderInfo.length}} Mermaid placeholders.`);

                                        // Load Mermaid
                                        if (!window.mermaid) {{
                                            await new Promise((resolve, reject) => {{
                                                const script = document.createElement("script");
                                                script.src = "{self.valves.MERMAID_JS_URL}";
                                                script.onload = resolve;
                                                script.onerror = reject;
                                                document.head.appendChild(script);
                                            }});
                                        }}

                                        mermaid.initialize({{
                                            startOnLoad: false,
                                            theme: 'default',
                                        }});

                                        // Read rels XML once
                                        const relsXml = await zip.file("word/_rels/document.xml.rels").async("string");
                                        const relsDoc = parser.parseFromString(relsXml, "application/xml");
                                        const relationships = relsDoc.getElementsByTagName("Relationship");
                                        const rIdToPath = {{}};
                                        
                                        for (let i = 0; i < relationships.length; i++) {{
                                            const rel = relationships[i];
                                            rIdToPath[rel.getAttribute("Id")] = rel.getAttribute("Target");
                                        }}

                                        // Render and replace
                                        console.log(`Processing ${{placeholderInfo.length}} diagrams...`);
                                        
                                        for (let i = 0; i < placeholderInfo.length; i++) {{
                                            const {{ rId, extent, code }} = placeholderInfo[i];
                                            const imagePath = "word/" + rIdToPath[rId];
                                            
                                            console.log(`Block ${{i + 1}}/${{placeholderInfo.length}}: Rendering and replacing at ${{imagePath}}`);

                                            // Render SVG
                                            const id = "mermaid-export-" + i;
                                            const {{ svg }} = await mermaid.render(id, code);
                                            
                                            // Convert SVG to PNG
                                            const canvas = document.createElement("canvas");
                                            const ctx = canvas.getContext("2d");
                                            const img = new Image();
                                            
                                            // Get SVG dimensions
                                            const svgMatch = svg.match(/viewBox="[^"]*\s+[^"]*\s+([^"\s]+)\s+([^"\s]+)"/);
                                            let width = 800;
                                            let height = 600;
                                            if (svgMatch) {{
                                                width = parseFloat(svgMatch[1]);
                                                height = parseFloat(svgMatch[2]);
                                            }}

                                            // Scale up for better quality
                                            const scale = {self.valves.MERMAID_PNG_SCALE};
                                            canvas.width = width * scale;
                                            canvas.height = height * scale;
                                            
                                            await new Promise((resolve, reject) => {{
                                                img.onload = resolve;
                                                img.onerror = reject;
                                                img.src = "data:image/svg+xml;base64," + btoa(unescape(encodeURIComponent(svg)));
                                            }});

                                            ctx.scale(scale, scale);
                                            ctx.drawImage(img, 0, 0, width, height);
                                            
                                            const pngDataUrl = canvas.toDataURL("image/png");
                                            const pngBase64 = pngDataUrl.split(",")[1];

                                            // Replace image in ZIP
                                            zip.file(imagePath, pngBase64, {{base64: true}});
                                            
                                            // Update dimensions in document.xml (EMUs)
                                            // 1 inch = 914400 EMUs, 1 pixel ≈ 9525 EMUs at 96 DPI
                                            // Max width: ~6 inches (page width minus margins)
                                            const maxWidthEmu = 5486400; // 6 inches
                                            const displayScale = {self.valves.MERMAID_DISPLAY_SCALE};
                                            let emuWidth = Math.round(width * 9525 * displayScale);
                                            let emuHeight = Math.round(height * 9525 * displayScale);
                                            
                                            // Scale down if too wide
                                            if (emuWidth > maxWidthEmu) {{
                                                const scaleFactor = maxWidthEmu / emuWidth;
                                                emuWidth = maxWidthEmu;
                                                emuHeight = Math.round(emuHeight * scaleFactor);
                                            }}
                                            
                                            extent.setAttribute("cx", emuWidth);
                                            extent.setAttribute("cy", emuHeight);
                                        }}

                                        // Serialize updated XML
                                        const serializer = new XMLSerializer();
                                        const newDocXml = serializer.serializeToString(xmlDoc);
                                        zip.file("word/document.xml", newDocXml);

                                        // Generate final blob
                                        const finalBlob = await zip.generateAsync({{type: "blob"}});
                                        const filename = "{filename}";

                                        const url = URL.createObjectURL(finalBlob);
                                        const a = document.createElement("a");
                                        a.style.display = "none";
                                        a.href = url;
                                        a.download = filename;
                                        document.body.appendChild(a);
                                        a.click();
                                        URL.revokeObjectURL(url);
                                        document.body.removeChild(a);
                                    }} catch (error) {{
                                        console.error('Export failed:', error);
                                        alert('导出失败: ' + error.message);
                                    }}
                                }})();
                                """
                            },
                        }
                    )

                await __event_emitter__(
                    {
                        "type": "status",
                        "data": {"description": "Word 文档已导出", "done": True},
                    }
                )

                await self._send_notification(
                    __event_emitter__, "success", f"已成功导出为 {filename}"
                )

                return {"message": "下载事件已触发"}

            except Exception as e:
                print(f"Error exporting to Word: {str(e)}")
                await __event_emitter__(
                    {
                        "type": "status",
                        "data": {
                            "description": f"导出失败: {str(e)}",
                            "done": True,
                        },
                    }
                )
                await self._send_notification(
                    __event_emitter__, "error", f"导出 Word 文档时出错: {str(e)}"
                )

    async def generate_title_using_ai(
        self, body: dict, content: str, user_id: str, request: Any
    ) -> str:
        if not request:
            return ""

        try:
            user_obj = Users.get_user_by_id(user_id)
            model = body.get("model")

            payload = {
                "model": model,
                "messages": [
                    {
                        "role": "system",
                        "content": "You are a helpful assistant. Generate a short, concise title (max 10 words) for the following text. Do not use quotes. Only output the title.",
                    },
                    {"role": "user", "content": content[:2000]},  # Limit content length
                ],
                "stream": False,
            }

            response = await generate_chat_completion(request, payload, user_obj)
            if response and "choices" in response:
                return response["choices"][0]["message"]["content"].strip()
        except Exception as e:
            logger.error(f"Error generating title: {e}")

        return ""

    def extract_title(self, content: str) -> str:
        """从 Markdown 内容提取一级/二级标题"""
        lines = content.split("\n")
        for line in lines:
            # 仅匹配 h1-h2 标题
            match = re.match(r"^#{1,2}\s+(.+)$", line.strip())
            if match:
                return match.group(1).strip()
        return ""

    def extract_chat_title(self, body: dict) -> str:
        """从请求体中提取会话标题"""
        if not isinstance(body, dict):
            return ""

        candidates = []

        for key in ("chat", "conversation"):
            if isinstance(body.get(key), dict):
                candidates.append(body.get(key, {}).get("title", ""))

        for key in ("title", "chat_title"):
            value = body.get(key)
            if isinstance(value, str):
                candidates.append(value)

        for candidate in candidates:
            if candidate and isinstance(candidate, str):
                return candidate.strip()
        return ""

    def extract_chat_id(self, body: dict, metadata: Optional[dict]) -> str:
        """从 body 或 metadata 中提取 chat_id"""
        if isinstance(body, dict):
            chat_id = body.get("chat_id") or body.get("id")
            if isinstance(chat_id, str) and chat_id.strip():
                return chat_id.strip()

            for key in ("chat", "conversation"):
                nested = body.get(key)
                if isinstance(nested, dict):
                    nested_id = nested.get("id") or nested.get("chat_id")
                    if isinstance(nested_id, str) and nested_id.strip():
                        return nested_id.strip()
        if isinstance(metadata, dict):
            chat_id = metadata.get("chat_id")
            if isinstance(chat_id, str) and chat_id.strip():
                return chat_id.strip()
        return ""

    async def fetch_chat_title(self, chat_id: str, user_id: str = "") -> str:
        """根据 chat_id 从数据库获取标题"""
        if not chat_id:
            return ""

        def _load_chat():
            if user_id:
                return Chats.get_chat_by_id_and_user_id(id=chat_id, user_id=user_id)
            return Chats.get_chat_by_id(chat_id)

        try:
            chat = await asyncio.to_thread(_load_chat)
        except Exception as exc:
            logger.warning(f"加载聊天 {chat_id} 失败: {exc}")
            return ""

        if not chat:
            return ""

        data = getattr(chat, "chat", {}) or {}
        title = data.get("title") or getattr(chat, "title", "")
        return title.strip() if isinstance(title, str) else ""

    def clean_filename(self, name: str) -> str:
        """清理文件名中的非法字符"""
        return re.sub(r'[\\/*?:"<>|]', "", name).strip()[:50]

    def _strip_reasoning_blocks(self, text: str) -> str:
        """
        Strip model reasoning blocks from assistant Markdown before export.
        """
        if not text:
            return text

        cur = text
        for _ in range(10):
            prev = cur
            cur = _REASONING_DETAILS_RE.sub("", cur)
            cur = _THINK_RE.sub("", cur)
            cur = _ANALYSIS_RE.sub("", cur)
            if cur == prev:
                break

        # Clean up excessive blank lines left by removals.
        cur = re.sub(r"\n{4,}", "\n\n\n", cur)
        return cur

    def _add_display_equation(self, doc: Document, latex: str):
        latex = (latex or "").strip()
        if not latex:
            return

        if not LATEX_MATH_AVAILABLE:
            self.add_code_block(doc, latex, "latex")
            return

        try:
            mathml = latex_to_mathml(latex)
            omml = mathml2omml.convert(mathml)
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cast(Any, para)._p.append(self._wrap_omml_for_word(omml))
        except Exception as exc:
            logger.warning(f"Math conversion failed; falling back to text: {exc}")
            self.add_code_block(doc, latex, "latex")

    def _wrap_omml_for_word(self, omml: str):
        m_ns = "http://schemas.openxmlformats.org/officeDocument/2006/math"
        w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        # Keep the OMML payload as-is, but ensure it has the math namespace declared.
        xml = f'<m:oMathPara xmlns:m="{m_ns}" xmlns:w="{w_ns}">{omml}</m:oMathPara>'
        return parse_xml(xml)

    def _add_inline_equation(
        self,
        paragraph,
        latex: str,
        bold: bool = False,
        italic: bool = False,
        strike: bool = False,
    ):
        latex = (latex or "").strip()
        if not latex:
            return

        if not LATEX_MATH_AVAILABLE:
            self._add_text_run(
                paragraph, f"\\({latex}\\)", bold=bold, italic=italic, strike=strike
            )
            return

        try:
            mathml = latex_to_mathml(latex)
            omml = mathml2omml.convert(mathml)
            o_math = self._omml_oMath_element(omml)
            run = paragraph.add_run()
            run.bold = bold
            run.italic = italic
            run.font.strike = strike
            cast(Any, run)._r.append(o_math)
        except Exception as exc:
            logger.warning(f"Inline math conversion failed; keeping literal: {exc}")
            self._add_text_run(
                paragraph, f"\\({latex}\\)", bold=bold, italic=italic, strike=strike
            )

    def _omml_oMath_element(self, omml: str):
        # Ensure the OMML element declares the math namespace so parse_xml works.
        m_ns = "http://schemas.openxmlformats.org/officeDocument/2006/math"
        s = (omml or "").strip()
        if s.startswith("<m:oMath>") and s.endswith("</m:oMath>"):
            inner = s[len("<m:oMath>") : -len("</m:oMath>")]
            s = f'<m:oMath xmlns:m="{m_ns}">{inner}</m:oMath>'
        elif s.startswith("<m:oMath") and "xmlns:m=" not in s.split(">", 1)[0]:
            s = s.replace("<m:oMath", f'<m:oMath xmlns:m="{m_ns}"', 1)
        return parse_xml(s)

    def _build_citation_refs(self, sources: List[dict]) -> List[_CitationRef]:
        citation_idx_map: Dict[str, int] = {}
        refs_by_idx: Dict[int, _CitationRef] = {}

        for source in sources or []:
            if not isinstance(source, dict):
                continue

            documents = source.get("document") or []
            metadatas = source.get("metadata") or []
            src_info = source.get("source") or {}

            src_name = src_info.get("name") if isinstance(src_info, dict) else None
            src_id_default = src_info.get("id") if isinstance(src_info, dict) else None
            src_urls = src_info.get("urls") if isinstance(src_info, dict) else None

            if not isinstance(documents, list):
                documents = []
            if not isinstance(metadatas, list):
                metadatas = []

            for idx_doc, _doc_text in enumerate(documents):
                meta = metadatas[idx_doc] if idx_doc < len(metadatas) else {}
                if not isinstance(meta, dict):
                    meta = {}

                source_id = meta.get("source") or src_id_default or "N/A"
                source_id_str = str(source_id)

                if source_id_str not in citation_idx_map:
                    citation_idx_map[source_id_str] = len(citation_idx_map) + 1
                idx = citation_idx_map[source_id_str]

                if idx in refs_by_idx:
                    continue

                url: Optional[str] = None
                if isinstance(source_id, str) and re.match(r"^https?://", source_id):
                    url = source_id
                elif isinstance(meta.get("url"), str) and re.match(
                    r"^https?://", meta["url"]
                ):
                    url = meta["url"]
                elif isinstance(src_urls, list) and src_urls:
                    if isinstance(src_urls[0], str) and re.match(
                        r"^https?://", src_urls[0]
                    ):
                        url = src_urls[0]

                title = (
                    (meta.get("title") if isinstance(meta.get("title"), str) else None)
                    or (meta.get("name") if isinstance(meta.get("name"), str) else None)
                    or (
                        src_name
                        if isinstance(src_name, str) and src_name.strip()
                        else None
                    )
                    or (url if url else None)
                    or source_id_str
                )

                anchor = f"OWUIRef{idx}"
                refs_by_idx[idx] = _CitationRef(
                    idx=idx,
                    anchor=anchor,
                    title=title,
                    url=url,
                    source_id=source_id_str,
                )

        return [refs_by_idx[i] for i in sorted(refs_by_idx.keys())]

    def _add_bookmark(self, paragraph, name: str):
        bookmark_id = self._bookmark_id_counter
        self._bookmark_id_counter += 1

        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(bookmark_id))
        start.set(qn("w:name"), name)

        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(bookmark_id))

        p = cast(Any, paragraph)._p
        p.insert(0, start)
        p.append(end)

    def _add_internal_hyperlink(self, paragraph, display_text: str, anchor: str):
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), anchor)

        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)

        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.text = display_text
        new_run.append(t)

        hyperlink.append(new_run)
        cast(Any, paragraph)._p.append(hyperlink)

    def _add_references_section(self, doc: Document):
        self.add_heading(doc, "参考资料", 2)

        for ref in self._citation_refs:
            para = doc.add_paragraph(style="List Number")
            self._add_bookmark(para, ref.anchor)
            # Include URL as an external link when available.
            if ref.url:
                self._add_hyperlink(para, ref.title, ref.url, display_text=ref.title)
            else:
                self._add_text_run(
                    para, ref.title, bold=False, italic=False, strike=False
                )

    def _normalize_url(self, url: str) -> str:
        u = (url or "").strip()
        if u.lower().startswith("www."):
            u = "https://" + u

        # Trim common trailing punctuation that often follows URLs in prose.
        while u and u[-1] in ".,;:!?)]}":
            u = u[:-1]
        return u

    def _add_hyperlink(
        self, paragraph, text: str, url: str, display_text: Optional[str] = None
    ):
        u = self._normalize_url(url)
        if not u:
            paragraph.add_run(display_text or text)
            return

        part = getattr(paragraph, "part", None)
        if part is None or not hasattr(part, "relate_to"):
            # Fallback if relationship API isn't available.
            run = paragraph.add_run(display_text or text)
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.font.underline = True
            return

        r_id = part.relate_to(u, RT.HYPERLINK, is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)

        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0000FF")
        rPr.append(color)

        u_el = OxmlElement("w:u")
        u_el.set(qn("w:val"), "single")
        rPr.append(u_el)

        new_run.append(rPr)

        t = OxmlElement("w:t")
        t.text = display_text or text
        new_run.append(t)

        hyperlink.append(new_run)
        cast(Any, paragraph)._p.append(hyperlink)

    def _add_text_run(self, paragraph, s: str, bold: bool, italic: bool, strike: bool):
        if not s:
            return
        run = paragraph.add_run(s)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if strike:
            run.font.strike = True

        # Set Chinese font (copying from existing add_paragraph logic)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def _add_inline_code(self, paragraph, s: str):
        if s == "":
            return

        # Simple inline code without URL parsing for now, or copy full logic if needed.
        # For now, just basic styling to match existing.
        run = paragraph.add_run(s)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
        run.font.size = Pt(10)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "E8E8E8")
        run._element.rPr.append(shading)

    def _add_hyperlink_code(self, paragraph, display_text: str, url: str):
        u = self._normalize_url(url)
        if not u:
            self._add_inline_code(paragraph, display_text)
            return

        part = getattr(paragraph, "part", None)
        if part is None or not hasattr(part, "relate_to"):
            self._add_inline_code(paragraph, display_text)
            return

        r_id = part.relate_to(u, RT.HYPERLINK, is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Consolas")
        rFonts.set(qn("w:hAnsi"), "Consolas")
        rFonts.set(qn("w:eastAsia"), "SimHei")
        rPr.append(rFonts)

        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "20")  # 10pt
        rPr.append(sz)

        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "E8E8E8")
        rPr.append(shading)

        new_run.append(rPr)

        t = OxmlElement("w:t")
        t.text = display_text
        new_run.append(t)

        hyperlink.append(new_run)
        cast(Any, paragraph)._p.append(hyperlink)

    def markdown_to_docx(
        self,
        markdown_text: str,
        top_heading: str = "",
        has_h1: bool = False,
        sources: Optional[List[dict]] = None,
    ) -> Document:
        """
        将 Markdown 文本转换为 Word 文档
        支持：标题、段落、粗体、斜体、代码块、列表、表格、链接、原生数学公式、引用和移除思考过程。
        """
        doc = Document()

        # 设置默认中文字体
        self.set_document_default_font(doc)

        # 构建引用
        self._citation_refs = self._build_citation_refs(sources)

        # 移除思考过程
        markdown_text = self._strip_reasoning_blocks(markdown_text)

        # 若正文无一级标题且有对话标题，则作为一级标题写入
        if top_heading and not has_h1:
            self.add_heading(doc, top_heading, 1)

        lines = markdown_text.split("\n")
        i = 0
        in_code_block = False
        code_block_content = []
        code_block_lang = ""
        in_list = False
        list_items = []
        list_type = None  # 'ordered' or 'unordered'

        while i < len(lines):
            line = lines[i]

            # 处理代码块
            if line.strip().startswith("```"):
                if not in_code_block:
                    # 先处理之前积累的列表
                    if in_list and list_items:
                        self.add_list_to_doc(doc, list_items, list_type)
                        list_items = []
                        in_list = False

                    in_code_block = True
                    code_block_info_raw = line.strip()[3:].strip()
                    code_block_lang, code_block_attrs = self._parse_fence_info(
                        code_block_info_raw
                    )
                    code_block_content = []
                else:
                    # 代码块结束
                    in_code_block = False
                    code_text = "\n".join(code_block_content)

                    # 检查是否为 Mermaid 或 Flowchart
                    mermaid_langs = {
                        "mermaid",
                        "flowchart",
                        "sequence",
                        "gantt",
                        "class",
                        "state",
                        "pie",
                        "er",
                        "journey",
                        "gitgraph",
                        "mindmap",
                    }

                    if code_block_lang.lower() in mermaid_langs:
                        # 创建 Mermaid 块对象
                        block = _MermaidFenceBlock(
                            info_raw=code_block_info_raw,
                            language=code_block_lang,
                            attrs=code_block_attrs,
                            source=code_text,
                        )
                        # 插入占位符
                        self._insert_mermaid_placeholder(doc, block)
                    else:
                        self.add_code_block(doc, code_text, code_block_lang)

                    code_block_content = []
                    code_block_lang = ""
                i += 1
                continue

            if in_code_block:
                code_block_content.append(line)
                i += 1
                continue

            # 处理数学块: $$...$$ 或 \[...\]
            # 简单检测: 如果行以 $$ 或 \[ 开头，则视为数学块开始
            stripped_line = line.strip()
            if stripped_line.startswith("$$") or stripped_line.startswith("\\["):
                # 先处理之前积累的列表
                if in_list and list_items:
                    self.add_list_to_doc(doc, list_items, list_type)
                    list_items = []
                    in_list = False

                # 检查是否为单行块，如 $$ E=mc^2 $$
                if (
                    stripped_line.startswith("$$")
                    and stripped_line.endswith("$$")
                    and len(stripped_line) > 4
                ) or (
                    stripped_line.startswith("\\[")
                    and stripped_line.endswith("\\]")
                    and len(stripped_line) > 4
                ):
                    # 提取内容
                    if stripped_line.startswith("$$"):
                        math_content = stripped_line[2:-2]
                    else:
                        math_content = stripped_line[2:-2]
                    self._add_display_equation(doc, math_content)
                    i += 1
                    continue

                # 多行数学块
                math_lines = []
                # 移除开头标记
                if stripped_line.startswith("$$"):
                    current_line_content = stripped_line[2:]
                    end_marker = "$$"
                else:
                    current_line_content = stripped_line[2:]
                    end_marker = "\\]"

                if current_line_content.strip():
                    math_lines.append(current_line_content)

                i += 1
                block_closed = False
                while i < len(lines):
                    next_line = lines[i]
                    if next_line.strip().endswith(end_marker):
                        # 找到结束标记
                        content_before_end = next_line.strip()[: -len(end_marker)]
                        if content_before_end.strip():
                            math_lines.append(content_before_end)
                        block_closed = True
                        i += 1
                        break
                    math_lines.append(next_line)
                    i += 1

                self._add_display_equation(doc, "\n".join(math_lines))
                continue

            # 处理表格
            if line.strip().startswith("|") and line.strip().endswith("|"):
                # 先处理之前积累的列表
                if in_list and list_items:
                    self.add_list_to_doc(doc, list_items, list_type)
                    list_items = []
                    in_list = False

                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i])
                    i += 1
                self.add_table(doc, table_lines)
                continue

            # 处理标题
            header_match = re.match(r"^(#{1,6})\s+(.+)$", line.strip())
            if header_match:
                # 先处理之前积累的列表
                if in_list and list_items:
                    self.add_list_to_doc(doc, list_items, list_type)
                    list_items = []
                    in_list = False

                level = len(header_match.group(1))
                text = header_match.group(2)
                self.add_heading(doc, text, level)
                i += 1
                continue

            # 处理无序列表
            unordered_match = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
            if unordered_match:
                if not in_list or list_type != "unordered":
                    if in_list and list_items:
                        self.add_list_to_doc(doc, list_items, list_type)
                        list_items = []
                    in_list = True
                    list_type = "unordered"
                indent = len(unordered_match.group(1)) // 2
                list_items.append((indent, unordered_match.group(2)))
                i += 1
                continue

            # 处理有序列表
            ordered_match = re.match(r"^(\s*)\d+[.)]\s+(.+)$", line)
            if ordered_match:
                if not in_list or list_type != "ordered":
                    if in_list and list_items:
                        self.add_list_to_doc(doc, list_items, list_type)
                        list_items = []
                    in_list = True
                    list_type = "ordered"
                indent = len(ordered_match.group(1)) // 2
                list_items.append((indent, ordered_match.group(2)))
                i += 1
                continue

            # 处理引用块
            if line.strip().startswith(">"):
                # 先处理之前积累的列表
                if in_list and list_items:
                    self.add_list_to_doc(doc, list_items, list_type)
                    list_items = []
                    in_list = False

                # 收集连续的引用行
                blockquote_lines = []
                while i < len(lines) and lines[i].strip().startswith(">"):
                    # 移除开头的 > 和可能的空格
                    quote_line = re.sub(r"^>\s?", "", lines[i])
                    blockquote_lines.append(quote_line)
                    i += 1
                self.add_blockquote(doc, "\n".join(blockquote_lines))
                continue

            # 处理水平分割线
            if re.match(r"^[-*_]{3,}$", line.strip()):
                # 先处理之前积累的列表
                if in_list and list_items:
                    self.add_list_to_doc(doc, list_items, list_type)
                    list_items = []
                    in_list = False

                self.add_horizontal_rule(doc)
                i += 1
                continue

            # 处理空行
            if not line.strip():
                # 列表结束
                if in_list and list_items:
                    self.add_list_to_doc(doc, list_items, list_type)
                    list_items = []
                    in_list = False
                i += 1
                continue

            # 处理普通段落
            if in_list and list_items:
                self.add_list_to_doc(doc, list_items, list_type)
                list_items = []
                in_list = False

            self.add_paragraph(doc, line)
            i += 1

        # 处理剩余的列表
        if in_list and list_items:
            self.add_list_to_doc(doc, list_items, list_type)

        # 添加参考资料章节
        if self._citation_refs:
            self._add_references_section(doc)

        return doc

    def set_document_default_font(self, doc: Document):
        """设置文档默认字体，确保中英文都正常显示"""
        # 设置正文样式
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"  # 英文字体
        font.size = Pt(11)

        # 设置中文字体
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        # 设置段落格式
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph_format.space_after = Pt(6)

    def add_heading(self, doc: Document, text: str, level: int):
        """添加标题"""
        # Word 标题级别从 0 开始，Markdown 从 1 开始
        heading_level = min(level, 9)  # Word 最多支持 Heading 9
        heading = doc.add_heading(level=heading_level)

        # 解析并添加格式化文本
        self.add_formatted_text(heading, text)

        # 设置中文字体
        for run in heading.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            run.font.color.rgb = RGBColor(0, 0, 0)

    def add_paragraph(self, doc: Document, text: str):
        """添加段落，支持内联格式"""
        paragraph = doc.add_paragraph()
        self.add_formatted_text(paragraph, text)

        # 设置中文字体
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def add_formatted_text(self, paragraph, text: str):
        """
        解析 Markdown 内联格式并添加到段落
        支持：粗体、斜体、行内代码、链接、删除线、行内公式、引用
        """
        # 定义格式化模式
        patterns = [
            # Inline Math \( ... \)
            (r"\\\((.+?)\\\)", {"math": True}),
            # Inline Math $...$ (single dollar signs, non-greedy)
            (r"(?<!\$)\$(?!\$)([^$]+?)\$(?!\$)", {"math": True}),
            # Citations [1], [2], etc.
            (r"\[(\d+)\]", {"citation": True}),
            # 粗斜体 ***text*** 或 ___text___
            (r"\*\*\*(.+?)\*\*\*|___(.+?)___", {"bold": True, "italic": True}),
            # 粗体 **text** 或 __text__
            (r"\*\*(.+?)\*\*|__(.+?)__", {"bold": True}),
            # 斜体 *text* 或 _text_
            (
                r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)|(?<!_)_(?!_)(.+?)(?<!_)_(?!_)",
                {"italic": True},
            ),
            # 行内代码 `code`
            (r"`([^`]+)`", {"code": True}),
            # 链接 [text](url)
            (r"\[([^\]]+)\]\(([^)]+)\)", {"link": True}),
            # 删除线 ~~text~~
            (r"~~(.+?)~~", {"strike": True}),
        ]

        # 收集所有匹配项
        all_matches = []

        for pattern, style in patterns:
            for match in re.finditer(pattern, text):
                # 获取匹配的文本内容
                groups = match.groups()
                matched_text = next((g for g in groups if g is not None), "")

                # Special handling for citations to ensure they map to valid refs
                if style.get("citation"):
                    try:
                        idx = int(matched_text)
                        # Only treat as citation if we have a corresponding reference
                        # Check if idx exists in our refs list (1-based index)
                        if not any(r.idx == idx for r in self._citation_refs):
                            continue
                    except ValueError:
                        continue

                all_matches.append(
                    {
                        "start": match.start(),
                        "end": match.end(),
                        "text": matched_text,
                        "style": style,
                        "full_match": match.group(0),
                        "url": (
                            groups[1] if style.get("link") and len(groups) > 1 else None
                        ),
                    }
                )

        # 按位置排序
        all_matches.sort(key=lambda x: x["start"])

        # 移除重叠的匹配项
        filtered_matches = []
        last_end = 0
        for m in all_matches:
            if m["start"] >= last_end:
                filtered_matches.append(m)
                last_end = m["end"]

        # 构建最终文本
        pos = 0
        for match in filtered_matches:
            # 添加匹配项之前的普通文本
            if match["start"] > pos:
                plain_text = text[pos : match["start"]]
                if plain_text:
                    self._add_text_run(paragraph, plain_text, False, False, False)

            # 添加格式化文本
            style = match["style"]
            run_text = match["text"]

            if style.get("math"):
                self._add_inline_equation(paragraph, run_text)
            elif style.get("citation"):
                idx = int(run_text)
                # Find the anchor for this index
                ref = next((r for r in self._citation_refs if r.idx == idx), None)
                if ref:
                    self._add_internal_hyperlink(paragraph, f"[{idx}]", ref.anchor)
                else:
                    self._add_text_run(paragraph, f"[{idx}]", False, False, False)
            elif style.get("link"):
                # 处理链接
                self._add_hyperlink(paragraph, run_text, match["url"])
            elif style.get("code"):
                # 行内代码
                self._add_inline_code(paragraph, run_text)
            else:
                # For bold/italic/strike, check if the text contains inline math
                # Pattern for inline math: \(...\) or $...$
                math_pattern = r"(\\\((.+?)\\\)|\$([^$]+?)\$)"
                math_matches = list(re.finditer(math_pattern, run_text))

                if math_matches:
                    # Process text with inline math
                    text_pos = 0
                    for math_match in math_matches:
                        # Add text before math
                        if math_match.start() > text_pos:
                            before_text = run_text[text_pos : math_match.start()]
                            self._add_text_run(
                                paragraph,
                                before_text,
                                bold=style.get("bold", False),
                                italic=style.get("italic", False),
                                strike=style.get("strike", False),
                            )
                        # Add inline equation with formatting
                        latex_content = math_match.group(2) or math_match.group(3)
                        self._add_inline_equation(
                            paragraph,
                            latex_content,
                            bold=style.get("bold", False),
                            italic=style.get("italic", False),
                            strike=style.get("strike", False),
                        )
                        text_pos = math_match.end()
                    # Add remaining text after last math
                    if text_pos < len(run_text):
                        self._add_text_run(
                            paragraph,
                            run_text[text_pos:],
                            bold=style.get("bold", False),
                            italic=style.get("italic", False),
                            strike=style.get("strike", False),
                        )
                else:
                    self._add_text_run(
                        paragraph,
                        run_text,
                        bold=style.get("bold", False),
                        italic=style.get("italic", False),
                        strike=style.get("strike", False),
                    )

            pos = match["end"]

        if pos < len(text):
            self._add_text_run(paragraph, text[pos:], False, False, False)

    def add_code_block(self, doc: Document, code: str, language: str = ""):
        """添加代码块，支持语法高亮"""
        # 语法高亮颜色映射 (基于常见的 IDE 配色)
        TOKEN_COLORS = {
            Token.Keyword: RGBColor(0, 92, 197),  # macOS 风格蓝 - 关键字
            Token.Keyword.Constant: RGBColor(0, 92, 197),
            Token.Keyword.Declaration: RGBColor(0, 92, 197),
            Token.Keyword.Namespace: RGBColor(0, 92, 197),
            Token.Keyword.Type: RGBColor(0, 92, 197),
            Token.Name.Function: RGBColor(0, 0, 0),  # 函数名保持黑色
            Token.Name.Class: RGBColor(38, 82, 120),  # 深青蓝 - 类名
            Token.Name.Decorator: RGBColor(170, 51, 0),  # 暖橙 - 装饰器
            Token.Name.Builtin: RGBColor(0, 110, 71),  # 墨绿 - 内置
            Token.String: RGBColor(196, 26, 22),  # 红色 - 字符串
            Token.String.Doc: RGBColor(109, 120, 133),  # 灰 - 文档字符串
            Token.Comment: RGBColor(109, 120, 133),  # 灰 - 注释
            Token.Comment.Single: RGBColor(109, 120, 133),
            Token.Comment.Multiline: RGBColor(109, 120, 133),
            Token.Number: RGBColor(28, 0, 207),  # 靛蓝 - 数字
            Token.Number.Integer: RGBColor(28, 0, 207),
            Token.Number.Float: RGBColor(28, 0, 207),
            Token.Operator: RGBColor(90, 99, 120),  # 灰蓝 - 运算符
            Token.Punctuation: RGBColor(0, 0, 0),  # 黑色 - 标点
        }

        def get_token_color(token_type):
            """递归查找 token 颜色"""
            while token_type:
                if token_type in TOKEN_COLORS:
                    return TOKEN_COLORS[token_type]
                token_type = token_type.parent
            return None

        # 添加语言标签（如果有）
        if language:
            lang_para = doc.add_paragraph()
            lang_para.paragraph_format.space_before = Pt(6)
            lang_para.paragraph_format.space_after = Pt(0)
            lang_para.paragraph_format.left_indent = Cm(0.5)
            lang_run = lang_para.add_run(language.upper())
            lang_run.font.name = "Consolas"
            lang_run.font.size = Pt(8)
            lang_run.font.color.rgb = RGBColor(100, 100, 100)
            lang_run.font.bold = True

        # 添加代码块段落
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.5)
        paragraph.paragraph_format.space_before = Pt(3) if language else Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)

        # 添加浅灰色背景
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F7F7F7")
        paragraph._element.pPr.append(shading)

        # 尝试使用 Pygments 进行语法高亮
        if PYGMENTS_AVAILABLE and language:
            try:
                lexer = get_lexer_by_name(language, stripall=False)
            except Exception:
                lexer = TextLexer()

            tokens = list(lex(code, lexer))

            for token_type, token_value in tokens:
                if not token_value:
                    continue
                run = paragraph.add_run(token_value)
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
                run.font.size = Pt(10)

                # 应用颜色
                color = get_token_color(token_type)
                if color:
                    run.font.color.rgb = color

                # 关键字加粗
                if token_type in Token.Keyword:
                    run.font.bold = True
        else:
            # 无语法高亮，纯文本显示
            run = paragraph.add_run(code)
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
            run.font.size = Pt(10)

    def _insert_mermaid_placeholder(
        self, doc: Document, block: Union[_MermaidFenceBlock, str]
    ):
        self._mermaid_figure_counter += 1
        if isinstance(block, str):
            code = block
        else:
            code = block.source

        # 为每个占位符创建唯一的透明 PNG
        # 通过改变图片尺寸来确保 python-docx 不会重用同一个图片文件
        # 使用 figure_counter 来创建不同尺寸（1x1, 1x2, 1x3, ...）
        from PIL import Image

        # 创建一个透明图片，尺寸为 1 x counter（确保每个都不同）
        img = Image.new("RGBA", (1, self._mermaid_figure_counter), (0, 0, 0, 0))
        image_stream = io.BytesIO()
        img.save(image_stream, format="PNG")
        image_stream.seek(0)

        # 添加居中段落
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()

        # 添加图片 (默认大小，稍后 JS 会调整)
        picture = run.add_picture(image_stream, width=Inches(1))

        # 设置 Alt Text (描述) 为 "MERMAID_SRC:<encoded_code>"
        # 这是 Python 和 JS 之间的魔法链接
        import urllib.parse

        encoded_code = urllib.parse.quote(code)

        # 访问底层 XML 设置 docPr descr
        inline = picture._inline
        docPr = inline.docPr

        # 使用 .set() 确保属性写入 XML
        docPr.set("descr", f"MERMAID_SRC:{encoded_code}")
        docPr.set("title", "Mermaid Diagram Placeholder")

        # 添加标题
        if self.valves.MERMAID_CAPTIONS_ENABLE:
            self._add_mermaid_caption(doc, self._mermaid_figure_counter)

    def _add_mermaid_caption(self, doc: Document, figure_number: int):
        if not self._caption_style_name:
            self._ensure_caption_style(doc)
            self._caption_style_name = self.valves.MERMAID_CAPTION_STYLE

        caption_text = f"{self.valves.MERMAID_CAPTION_PREFIX} {figure_number}"
        paragraph = doc.add_paragraph(caption_text, style=self._caption_style_name)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = False

    def _ensure_caption_style(self, doc: Document):
        style_name = self.valves.MERMAID_CAPTION_STYLE
        styles = doc.styles
        if style_name not in styles:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = styles["Normal"]
            style.next_paragraph_style = styles["Normal"]
            font = style.font
            font.name = "Times New Roman"
            font.size = Pt(10)
            font.italic = True
            font.color.rgb = RGBColor(0x55, 0x55, 0x55)  # 深灰色

    def _parse_fence_info(self, info_raw: str) -> Tuple[str, List[str]]:
        parts = info_raw.split()
        if not parts:
            return "", []
        lang = parts[0]
        attrs = parts[1:]
        return lang, attrs

    def add_table(self, doc: Document, table_lines: List[str]):
        """添加 Markdown 表格，支持智能列宽、对齐和单元格内格式化"""
        if len(table_lines) < 2:
            return

        header_fill = "F2F2F2"
        zebra_fill = "FBFBFB"

        def _split_row(line: str) -> List[str]:
            raw = line.strip().strip("|")
            return [c.strip() for c in raw.split("|")]

        def _is_separator_row(cells: List[str]) -> bool:
            if not cells:
                return False
            ok = 0
            for c in cells:
                c = c.strip()
                if re.fullmatch(r":?-{3,}:?", c):
                    ok += 1
            return ok == len(cells)

        def _col_align(cell: str) -> WD_ALIGN_PARAGRAPH:
            s = (cell or "").strip()
            if s.startswith(":") and s.endswith(":"):
                return WD_ALIGN_PARAGRAPH.CENTER
            if s.endswith(":"):
                return WD_ALIGN_PARAGRAPH.RIGHT
            return WD_ALIGN_PARAGRAPH.LEFT

        def _set_cell_shading(cell, fill: str):
            tc_pr = cell._element.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), fill)
            tc_pr.append(shd)

        raw_rows = [_split_row(l) for l in table_lines if l.strip().startswith("|")]
        if not raw_rows:
            return

        sep_idx = 1 if len(raw_rows) > 1 and _is_separator_row(raw_rows[1]) else -1
        header = raw_rows[0]
        body = raw_rows[sep_idx + 1 :] if sep_idx >= 0 else raw_rows[1:]

        num_cols = max(len(header), *(len(r) for r in body)) if body else len(header)
        header = header + [""] * (num_cols - len(header))
        body = [r + [""] * (num_cols - len(r)) for r in body]

        aligns = [
            _col_align(c) for c in (raw_rows[1] if sep_idx == 1 else [""] * num_cols)
        ]

        table = doc.add_table(rows=1 + len(body), cols=num_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        cast(Any, table).autofit = False

        # 单元格边距
        self._set_table_cell_margins(table, top=60, bottom=60, left=90, right=90)

        # 列宽按内容比例分配
        available_width = int(self._available_block_width(doc))
        min_col = max(int(Inches(0.55)), available_width // max(1, num_cols * 3))

        def _plain_len(s: str) -> int:
            t = re.sub(r"`([^`]+)`", r"\1", s or "")
            t = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", t)
            t = re.sub(r"\s+", " ", t).strip()
            return len(t)

        weights: List[int] = []
        for ci in range(num_cols):
            max_len = _plain_len(header[ci])
            for r in body:
                max_len = max(max_len, _plain_len(r[ci]))
            weights.append(max(1, min(max_len, 40)))

        sum_w = sum(weights) or 1
        widths = [max(min_col, int(available_width * w / sum_w)) for w in weights]
        total = sum(widths)
        if total > available_width:
            even = max(1, available_width // max(1, num_cols))
            widths = [even] * num_cols
            total = sum(widths)
        if total < available_width:
            rem = available_width - total
            order = sorted(range(num_cols), key=lambda i: weights[i], reverse=True)
            oi = 0
            while rem > 0 and order:
                widths[order[oi % len(order)]] += 1
                rem -= 1
                oi += 1

        for ci, w in enumerate(widths):
            table.columns[ci].width = w
            for row in table.rows:
                row.cells[ci].width = w

        def _format_cell_paragraph(para, align: WD_ALIGN_PARAGRAPH):
            para.alignment = align
            pf = para.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

        def _fill_cell(cell, text: str, align: WD_ALIGN_PARAGRAPH, bold: bool = False):
            cell.text = ""
            parts = [
                p for p in re.split(r"(?:<br\s*/?>|\\n)", text or "") if p is not None
            ]
            if not parts:
                parts = [""]
            for pi, part in enumerate(parts):
                para = cell.paragraphs[0] if pi == 0 else cell.add_paragraph()
                _format_cell_paragraph(para, align)
                self.add_formatted_text(para, part)
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    if bold:
                        run.bold = True

        # 表头行
        header_row = table.rows[0]
        self._set_table_header_row_repeat(header_row)
        for ci in range(num_cols):
            cell = header_row.cells[ci]
            _set_cell_shading(cell, header_fill)
            _fill_cell(
                cell,
                header[ci],
                aligns[ci] if ci < len(aligns) else WD_ALIGN_PARAGRAPH.LEFT,
                bold=True,
            )

        # 数据行
        for ri, row_data in enumerate(body, start=1):
            row = table.rows[ri]
            for ci in range(num_cols):
                cell = row.cells[ci]
                if (ri % 2) == 0:
                    _set_cell_shading(cell, zebra_fill)
                _fill_cell(
                    cell,
                    row_data[ci],
                    aligns[ci] if ci < len(aligns) else WD_ALIGN_PARAGRAPH.LEFT,
                )

    def _available_block_width(self, doc: Document):
        section = doc.sections[0]
        return section.page_width - section.left_margin - section.right_margin

    def _set_table_cell_margins(
        self, table, top: int, bottom: int, left: int, right: int
    ):
        tbl_pr = cast(Any, table)._tbl.tblPr
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        for tag, val in (
            ("top", top),
            ("bottom", bottom),
            ("left", left),
            ("right", right),
        ):
            el = OxmlElement(f"w:{tag}")
            el.set(qn("w:w"), str(int(val)))
            el.set(qn("w:type"), "dxa")
            tbl_cell_mar.append(el)
        tbl_pr.append(tbl_cell_mar)

    def _set_table_header_row_repeat(self, row):
        tr_pr = row._tr.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)

    def add_list_to_doc(
        self, doc: Document, items: List[Tuple[int, str]], list_type: str
    ):
        """添加列表"""
        for indent, text in items:
            paragraph = doc.add_paragraph()

            if list_type == "unordered":
                # 无序列表使用项目符号
                paragraph.style = "List Bullet"
            else:
                # 有序列表使用编号
                paragraph.style = "List Number"

            # 设置缩进
            paragraph.paragraph_format.left_indent = Cm(0.5 * (indent + 1))

            # 添加格式化文本
            self.add_formatted_text(paragraph, text)

            # 设置字体
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def add_horizontal_rule(self, doc: Document):
        """添加水平分割线"""
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)

        # 添加底部边框作为分割线
        pPr = paragraph._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_blockquote(self, doc: Document, text: str):
        """添加引用块，带有左侧边框和灰色背景"""
        for line in text.split("\n"):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(1.0)
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(3)

            # 添加左侧边框
            pPr = paragraph._element.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "24")  # 边框粗细
            left.set(qn("w:space"), "4")  # 边框与文字间距
            left.set(qn("w:color"), "CCCCCC")  # 灰色边框
            pBdr.append(left)
            pPr.append(pBdr)

            # 添加浅灰色背景
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F9F9F9")
            pPr.append(shading)

            # 添加格式化文本
            self.add_formatted_text(paragraph, line)

            # 设置字体为斜体灰色
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
                run.font.color.rgb = RGBColor(85, 85, 85)  # 深灰色文字
                run.italic = True
